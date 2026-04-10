"""
Extract plain text from uploaded appraisal files.
Supports: PDF, DOCX, TXT.
"""
import io
import os
import aiofiles
from pathlib import Path
from typing import Optional
import structlog

logger = structlog.get_logger()


async def extract_text_from_file(file_path: str, mime_type: str) -> str:
    """Read a saved file and extract its text content."""
    try:
        async with aiofiles.open(file_path, "rb") as f:
            content = await f.read()
        return extract_text_from_bytes(content, mime_type, file_path)
    except Exception as e:
        logger.error("file_extraction_failed", path=file_path, error=str(e))
        return ""


def extract_text_from_bytes(content: bytes, mime_type: str, filename: str = "") -> str:
    """Extract text from raw bytes based on MIME type."""
    fname_lower = filename.lower()

    # ── PDF ──────────────────────────────────────────────────────────────────
    if "pdf" in mime_type or fname_lower.endswith(".pdf"):
        return _extract_pdf(content)

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if "word" in mime_type or "officedocument" in mime_type or fname_lower.endswith(".docx"):
        return _extract_docx(content)

    # ── Plain text ────────────────────────────────────────────────────────────
    if "text" in mime_type or fname_lower.endswith(".txt"):
        return content.decode("utf-8", errors="replace")

    return content.decode("utf-8", errors="replace")


def _extract_pdf(content: bytes) -> str:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except Exception as e:
        logger.warning("pdf_extraction_failed", error=str(e))
        return ""


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.warning("docx_extraction_failed", error=str(e))
        return ""
